# 建立 Series
import pandas as pd
data = pd.Series([20, 10, 15])

# 基本操作
print("最大值:", data.max())
print("中位數:", data.median())
data = data * 2
print(data)
data = data == 20
print(data)


# 建立 DataFrame
data=pd.DataFrame({
    "name":["Amy", "John", "Bob"],
    "salary":[30000, 50000, 40000]
})
# 基本 DataFrame 操作
print(data)
#取得特定的欄位
print(data["salary"])
# 印出第一列
print(data.iloc[0])


# 創建 Series

import pandas as pd

# 使用列表創建 Series
data = [10, 20, 30, 40]
s = pd.Series(data)
print(s)

# 使用指定索引創建 Series
data = [10, 20, 30, 40]
index = ['a', 'b', 'c', 'd']
s = pd.Series(data, index=index)
print(s)


# Series 操作
# 1. 取出元素
# 取出索引為 'b' 的元素
print(s['b'])

# 取出位置為 2 的元素
print(s[2])

# 2. 更新元素
# 更新索引為 'a' 的元素
s['a'] = 100
print(s)


# 3. 基本統計
# 計算 Series 的總和
print(s.sum())

# 計算 Series 的均值
print(s.mean())

# 計算 Series 的標準差
print(s.std())

# 4. 過濾
# 過濾出大於 25 的數據
filtered_s = s[s > 25]
print(filtered_s)

# Series 與其他資料結構的轉換
# 1. Series 轉換為列表
# 將 Series 轉換為列表
list_data = s.tolist()
print(list_data)

# 2. Series 轉換為字典
# 將 Series 轉換為字典
dict_data = s.to_dict()
print(dict_data)

# Series 方法
# 查看 Series 的索引
print(s.index)

# 查看 Series 的數據
print(s.values)

# 查看 Series 的描述性統計信息
print(s.describe())


# 創建 DataFrame
import pandas as pd

# 使用字典創建 DataFrame
data = {
    'Name': ['Alice', 'Bob', 'Charlie'],
    'Age': [25, 30, 35],
    'City': ['New York', 'Los Angeles', 'Chicago']
}
df = pd.DataFrame(data)
print(df)

# 使用列表創建 DataFrame
data = [
    ['Alice', 25, 'New York'],
    ['Bob', 30, 'Los Angeles'],
    ['Charlie', 35, 'Chicago']
]
df = pd.DataFrame(data, columns=['Name', 'Age', 'City'])
print(df)

# DataFrame 操作
# 1. 取出行和列
# 取出特定的列
print(df['Name'])

# 取出特定的行（按位置）
print(df.iloc[1])

# 取出特定的行（按索引）
print(df.loc[1])

# 取出特定的行和列
print(df.loc[1, 'Name'])
print(df.iloc[1, 0])

# 2. 更新數據
# 更新指定行和列的數據
df.loc[1, 'Age'] = 31
print(df)

# 更新整列數據
df['City'] = ['San Francisco', 'San Diego', 'Houston']
print(df)

# 3. 基本統計
# 計算數字列的均值
print(df['Age'].mean())

# 計算數字列的總和
print(df['Age'].sum())

# 計算數字列的標準差
print(df['Age'].std())

# 4. 過濾數據
# 過濾出年齡大於 30 的行
filtered_df = df[df['Age'] > 30]
print(filtered_df)

# 5. 新增和刪除列
# 新增一列
df['Country'] = ['USA', 'USA', 'USA']
print(df)

# 刪除一列
df = df.drop('Country', axis=1)
print(df)

# DataFrame 與其他資料結構的轉換
# 1. DataFrame 轉換為字典
# 將 DataFrame 轉換為字典
dict_data = df.to_dict()
print(dict_data)

# 2. DataFrame 轉換為列表
# 將 DataFrame 轉換為列表
list_data = df.values.tolist()
print(list_data)

# DataFrame 方法
# 查看 DataFrame 的行和列
print(df.shape)

# 查看 DataFrame 的索引
print(df.index)

# 查看 DataFrame 的列名
print(df.columns)

# 查看 DataFrame 的描述性統計信息
print(df.describe())


import pandas as pd

# 創建 Series
print("創建 Series:")
# 使用列表創建 Series
data = [10, 20, 30, 40]
s = pd.Series(data)
print(s)

# 使用指定索引創建 Series
data = [10, 20, 30, 40]
index = ['a', 'b', 'c', 'd']
s = pd.Series(data, index=index)
print(s)

# Series 操作
print("\nSeries 操作:")
# 1. 取出元素
# 取出索引為 'b' 的元素
print("取出索引為 'b' 的元素:", s['b'])

# 取出位置為 2 的元素
print("取出位置為 2 的元素:", s[2])

# 2. 更新元素
# 更新索引為 'a' 的元素
s['a'] = 100
print("更新後的 Series:\n", s)

# 3. 基本統計
# 計算 Series 的總和
print("Series 的總和:", s.sum())

# 計算 Series 的均值
print("Series 的均值:", s.mean())

# 計算 Series 的標準差
print("Series 的標準差:", s.std())

# 4. 過濾
# 過濾出大於 25 的數據
filtered_s = s[s > 25]
print("過濾後的 Series:\n", filtered_s)

# Series 與其他資料結構的轉換
print("\nSeries 與其他資料結構的轉換:")
# 1. Series 轉換為列表
# 將 Series 轉換為列表
list_data = s.tolist()
print("Series 轉換為列表:", list_data)

# 2. Series 轉換為字典
# 將 Series 轉換為字典
dict_data = s.to_dict()
print("Series 轉換為字典:", dict_data)

# Series 方法
print("\nSeries 方法:")
# 查看 Series 的索引
print("Series 的索引:", s.index)

# 查看 Series 的數據
print("Series 的數據:", s.values)

# 查看 Series 的描述性統計信息
print("Series 的描述性統計信息:\n", s.describe())

# 創建 DataFrame
print("\n創建 DataFrame:")
# 使用字典創建 DataFrame
data = {
    'Name': ['Alice', 'Bob', 'Charlie'],
    'Age': [25, 30, 35],
    'City': ['New York', 'Los Angeles', 'Chicago']
}
df = pd.DataFrame(data)
print(df)

# 使用列表創建 DataFrame
data = [
    ['Alice', 25, 'New York'],
    ['Bob', 30, 'Los Angeles'],
    ['Charlie', 35, 'Chicago']
]
df = pd.Data




import mysql.connector
from mysql.connector import Error
conn=None
try:
    # 連接到 MySQL 資料庫
    conn = mysql.connector.connect(
        host="10.96.196.36",  # 伺服器 IP
        user="NPSPO_User",  # 使用者名稱
        password="Gaming@NPD",  # 密碼
        database="stan01"  # 資料庫名稱
    )
 
    if conn.is_connected():
        print("已成功連接到 MySQL 資料庫")
       
        # 創建游標對象
        cursor = conn.cursor()
 
        # 執行 SQL 查詢
        cursor.execute("SELECT * FROM speedtest_results LIMIT 10")  # speedtest_results 之後替換為你想查詢的資料表名稱
       
        # 獲取查詢結果
        results = cursor.fetchall()
       
        # 顯示查詢結果
        for row in results:
            print(row)
 
except Error as e:
    print(f"資料庫錯誤: {e}")
 
finally:
    # 確保連接在程式結束時正確關閉
    if conn.is_connected():
        cursor.close()
        conn.close()
        print("資料庫連接已關閉")

# 資料分析
# 建立Series資料
import pandas as pd
data=pd.Series('列表')
# 建立篩選條件(與資料數量對應的布林值)
condition=[True,False,True]
#根據條件完成篩選
filteredData=data[condition]

# 常見操作
# 建立Series資料
import pandas as pd
data=pd.Series('列表')
# 建立篩選條件(直接透過比較運算產生)
condition=data>5
#根據條件完成篩選
filteredData=data[condition]

# 基本邏輯
# 建立DataFrame資料
import pandas as pd
data=pd.DataFrame('字典')
# 建立篩選條件(透過特定攔位的比較運算的產生)
condition=data['欄位名稱']>5
#根據條件完成篩選
filteredData=data[condition]

# 載入pandas模組
import pandas as pd
# 篩選練習 - Series
data=pd.Series([30, 15, 20])
condition=data>18
print(condition)
filteredData=data[condition]
print(filteredData)

data=pd.Series(["您好", "Python", "Pandas"])
condition=data.str.contains("P")
print(condition)
filteredData=data[condition]
print(filteredData)

# 篩選練習 - DataFrame
data=pd.DataFrame({
    "name":["Amy", "Bob", "Charles"],
    "salary":[30000, 50000, 40000]
})
print(data)
condition=data["salary"]>=40000
print(condition)
filteredData=data[condition]
print(filteredData)

condition=data["name"]=="Amy"
print(condition)
filteredData=data[condition]
print(filteredData)

# 資料工程
# 清理資料
# 資料分析
# 基於資料的應用
# 收集資料
#清理、分析與應用
#下載檔案https://www.youtube.com/redirect?event=video_description&redir_token=QUFFLUhqbk00TmpZRFRXSVhyVTJoWHlzYWg2YUM3RU8tQXxBQ3Jtc0tuQnVNWUMtdFNsRlg0SS0ydkVfZmtjMG9GS3d6bjI5YnNEMFJlU1oxLWRRQVlUZnp6VDY1bDFOZGJTdnBNOU4zNktkUUNSTl9uN1FWYlJtU2x3UGhzZVM1MG0tV2IzNUJlc3I3U29JMzVhVDk0SW5Daw&q=https%3A%2F%2Fcwpeng.github.io%2Flive-records-samples%2Fdata%2Fgoogleplaystore.csv&v=B5BgPWBZhvY
import pandas as pd
# 讀取資料
data=pd.read_csv("googleplaystore.csv") # 把CSV格式檔案讀取成一個DataFrame
# 觀察資料
print(data)
print("資料數量", data.shape)
print("資料欄位", data.columns)
# 分析資料:評分的各種統計數據
condition=data["Rating"]<=5
data=data[condition]
print(data)
print( data["Rating"])
print("中位數", data["Rating"].mean())
print("平均數",data["Rating"].median())
print("取得前一千個應用成是的平均數", data["Rating"].nlargest(1000).mean())
# 分析資料:安裝數量的各種統計數據
print(data["Installs"])
print(data["Installs"][10472])

data["Installs"]=pd.to_numeric(data["Installs"].str.replace("[,+]",""))
data["Installs"] = pd.to_numeric(data["Installs"].str.replace("[,+]", "", regex=True))

print("平均數l", data["Installs"].mean())
# 基於資料的應用:關鍵字搜尋應用程式的名稱
keyword=input("請輸入關鍵字:")
condition=data["App"].str.contains(keyword, case=False)
print("包含關鍵字的應用程式數量", data[condition].shape[0])


# 資料工程
# 清理資料
# 資料分析
# 基於資料的應用
# 收集資料
#清理、分析與應用
#下載檔案https://www.youtube.com/redirect?event=video_description&redir_token=QUFFLUhqbk00TmpZRFRXSVhyVTJoWHlzYWg2YUM3RU8tQXxBQ3Jtc0tuQnVNWUMtdFNsRlg0SS0ydkVfZmtjMG9GS3d6bjI5YnNEMFJlU1oxLWRRQVlUZnp6VDY1bDFOZGJTdnBNOU4zNktkUUNSTl9uN1FWYlJtU2x3UGhzZVM1MG0tV2IzNUJlc3I3U29JMzVhVDk0SW5Daw&q=https%3A%2F%2Fcwpeng.github.io%2Flive-records-samples%2Fdata%2Fgoogleplaystore.csv&v=B5BgPWBZhvY
import pandas as pd

# 讀取資料
data = pd.read_csv("googleplaystore.csv")  # 把CSV格式檔案讀取成一個DataFrame

# 觀察資料
print(data)  # 打印出資料的內容
print("資料數量", data.shape)  # 顯示資料的行數和列數 (行數, 列數)
print("資料欄位", data.columns)  # 顯示資料的欄位名稱

# 分析資料: 評分的各種統計數據
# 篩選出評分小於等於 5 的資料，排除不合理的數據
condition = data["Rating"] <= 5
data = data[condition]  # 使用篩選條件過濾資料
print(data)  # 打印篩選後的資料
print(data["Rating"])  # 打印 "Rating" 欄位的數值

# 計算並打印 "Rating" 欄位的中位數和平均數
print("中位數", data["Rating"].median())  # 中位數
print("平均數", data["Rating"].mean())  # 平均數

# 取得 "Rating" 欄位前一千個最高評分，並計算它們的平均數
print("取得前一千個應用程式的平均數", data["Rating"].nlargest(1000).mean())

# 分析資料: 安裝數量的各種統計數據
print(data["Installs"])  # 打印 "Installs" 欄位的數值
print(data["Installs"][10472])  # 打印索引為 10472 的應用的安裝數量

# 清理 "Installs" 欄位中的數據，移除逗號和加號，並將其轉換為數值類型
data["Installs"] = pd.to_numeric(data["Installs"].str.replace("[,+]", "", regex=True))

# 計算並打印 "Installs" 欄位的平均數
print("平均數", data["Installs"].mean())

# 基於資料的應用: 關鍵字搜尋應用程式名稱
# 使用者輸入關鍵字，篩選應用程式名稱中包含關鍵字的資料
keyword = input("請輸入關鍵字:")  # 提示使用者輸入關鍵字
condition = data["App"].str.contains(keyword, case=False)  # 忽略大小寫，篩選出應用程式名稱中包含關鍵字的資料
print("包含關鍵字的應用程式數量", data[condition].shape[0])  # 打印符合條件的應用程式數量


# 連線網址

# 匯入 urllib.request 模組，用來發送網路請求
import urllib.request as request

# 發送一個網路請求，並開啟指定網址的內容
with request.urlopen("網址") as response:
    data = response.read()  # 讀取網站回傳的資料
print(data)  # 輸出資料 (通常是未處理過的原始資料)

# 取得網站的 HTML 原始碼
import urllib.request as request

src = "https://www.google.com"
with request.urlopen(src) as response:
    data = response.read().decode("utf-8")  # 取得網站的原始碼 (HTML、CSS、JS)
print(data)  # 輸出網站內容

# 從台北開放資料平台 API 取得 JSON 資料
import urllib.request as request
import json

src = "https://data.taipei/api/v1/dataset/296acfa2-5d93-4706-ad58-e83cc951863c?scope=resourceAquire"

# 發送請求並接收回應
with request.urlopen(src) as response:
      data = json.load(response)  # 解析 JSON 資料
print(data)
# 將公司名稱列印出來
clist = data["result"]["results"]
for company in clist :
    print(f"{company["公司名稱"]} : {company["公司地址"]}")

# 將結果寫入檔案
with open("data.txt","w", encoding="utf-8") as file:
    for company in clist:
        file.write(f"{company['公司名稱']} : {company['公司地址']}\n")  # 將公司名稱和地址寫入檔案
        # print(f"{company['公司名稱']} : {company['公司地址']}")  # 同時在螢幕上列出公司名稱和地址



# open 函數模式說明：
# "r"：以只讀模式打開檔案。如果檔案不存在，會引發 FileNotFoundError。
# "w"：以寫入模式打開檔案。如果檔案已存在，它會被覆蓋；如果檔案不存在，會創建一個新的檔案。
# "a"：以附加模式打開檔案。所有寫入操作都會追加到檔案末尾，而不是覆蓋原有內容。如果檔案不存在，會創建一個新的檔案。
# "b"：以二進位模式打開檔案（例如 "wb" 或 "rb"），用於處理二進位檔案如圖片或音頻文件。

# 讀取檔案內容
with open("file.txt", "r", encoding="utf-8") as file:
    content = file.read()
    print(content)

# 寫入檔案
with open("file.txt", "w", encoding="utf-8") as file:
    file.write("這是新內容\n")

# 附加檔案
with open("file.txt", "a", encoding="utf-8") as file:
    file.write("這是附加內容\n")

# 讀寫檔案
with open("file.txt", "r+", encoding="utf-8") as file:
    file.write("新的內容")
    file.seek(0)
    print(file.read())

# 二進位檔案操作
with open("image.jpg", "rb") as file:
    data = file.read()
    # 可以進行二進位數據處理


# 字符編碼（如 UTF-8）:

# 用途: 用於將文本數據轉換為二進制格式以便儲存，或將二進制數據轉換為文本格式以便讀取。
# 影響: 影響文本文件的讀取和寫入，確保文本中的字符（如中文、特殊符號）能夠正確顯示和處理。
# __________________________________________________________________

# 資料串接與讀寫

import pyodbc

conn_str = (
    "DRIVER={ODBC Driver 17 for SQL Server};"
    "SERVER=your_server;"
    "DATABASE=你的database;"
    "UID=your_username;"
    "PWD=your_password;"
)

# 建立連接

conn = pyodbc.connect(conn_str)
cursor = conn.cursor

#查詢資料
cursor.execute("SELECT * FROM your_table")
for row in cursor.fetchall():
    print(row)

# 關閉連結
cursor.close()
conn.close()


# 連接 MySQL（使用 mysql-connector-python）

import mysql.connector

# 建立連接
conn = mysql.connector.connect(
    host="your_host",
    user="your_username",
    password="your_password",
    database="yourdatabase"
)
cursor = conn.cursor()

# 查詢資料
cursor.execute("SELECT * FROM your_table")
for row in cursor.fetchall():
    print(row)

#關閉連接
cursor.close()
conn.close()


import mysql.connector
from mysql.connector import Error
conn=None
try:
    #連接到MySQL資料庫
    conn = mysql.connector.connect(
        host="10.96.196.36",
        user="NPSPD_User",
        password="Gaming@NPD",
        database="stan01"
    )
    
    if conn.is_connected():
        print("已成功連接到 MySQL資料庫")

        # 創建游標以便數據庫連接 執行SQL查詢或其他操作
        cursor = conn.cursor()

        # 執行SQL查詢
        cursor.execute("SELECT * FROM speedtest_results LIMIT 10") 
        # speedtest_results 之後替換為你想查詢的資料表名稱

        # 獲取查詢結果
        results = cursor.fetchall()

        # 顯示查詢結果
        for row in results:
            print(row)

except Error as e:
    print(f"資料庫錯誤: {e}")

finally:
    # 確保連接程式結束時正確關閉
    if conn.is_connected():
        cursor.close()
        conn.close()
        print("資料庫連接已關閉")

# 檔案的讀取與寫入

# CSV 檔案的讀取與寫入（使用 pandas）

import pandas as pd

# 讀取 CSV 檔案
df = pd.read_csv("file.csv")

# 處理資料
print(df.head())

# 寫入 CSV 檔案
df.to_csv('output.csv', index=False)

# index=False

# 假設你有一個 DataFrame：
import pandas as pd

data = {'Name': ['Alice', 'Bob'], 'Age': [25, 30]}
df = pd.DataFrame(data)
print(df)

# 使用 index=False：
df.to_csv('output.csv', index=False)

# output.csv 內容將是：

# Name,Age
# Alice,25
# Bob,30


import pandas as pd

# 讀取 Excel 檔案
df = pd.read_excel('file.xlsx', sheet_name='Sheet1')

# 處理資料
print(df.head())

# 寫入 Excel 檔按
df.to_excel('output.xlsx', index=False)

print(df)


# 直接創建資料Excel並自動開啟(如有更新內容請關閉後再次執行)
import pandas as pd
import openpyxl
import os

# 創建資料
data = {
    'Name': ['Lexi', 'Derek', 'Stan'],
    'Age': [18, 23, 22],
    'City': ['新店', '大安', '內湖']
}

# 將資料轉換為 DataFrame
df = pd.DataFrame(data)

# 將 DataFrame 寫入新的 Excel 檔案
df.to_excel('new_file.xlsx', index=False)

# 自動打開 Excel 檔案 (僅限於 Windows)
os.startfile(r'new_file.xlsx')

print("新的 Excel 檔案已建立並保存為 new_file.xlsx")

# Word 檔案的讀取與寫入（使用 python-docx）

from docx import Document

# 讀取 Word 文件
doc = Document('file.docx')

# 讀取每隔段落
for paragraph in doc.paragraphs:
    print(paragraph.text)

# 寫入 Word 文件
new_doc = Document()
new_doc.add_paragraph("這是一個新的段落")
new_doc.save('output.docx')

df = pd.DataFrame(data)

df.to_docx('new_word.docx', index=False)

os.startflie(r'new_word.docx')
print('新的 word 檔案已建立並保存為 new_word.docx')


import pandas as pd
from docx import Document
import os

# 假設你有一個資料字典或其他資料形式
data = {
    'Column1': [0, 9, 2, 8],  # 這裡的資料
    'Column2': ['S', 'T', 'A','N']  # 修改為對應數量的資料
}

# 將資料轉換為 DataFrame
df = pd.DataFrame(data)

# 創建新的 Word 文檔
new_doc = Document()

# 將 DataFrame 寫入 Word 文件
for index, row in df.iterrows():
    new_doc.add_paragraph(f"{row['Column1']} - {row['Column2']}")

new_doc.save('new_word.docx')

# 打開新的 Word 文件
os.startfile('new_word.docx')
print('新的 Word 檔案已建立並保存為 new_word.docx')

import pandas as pd
from docx import Document
import os

# 創建新的 Word 文檔
new_doc = Document()

# 添加兩行文字
new_doc.add_paragraph("910928")
new_doc.add_paragraph("Stan")

# 保存文檔
new_doc.save('new_word.docx')

# 打開新的 Word 文件
os.startfile('new_word.docx')
print('新的 Word 檔案已建立並保存為 new_word.docx')


# TXT檔案的讀取與寫入
# 基本語法:
# with open('filename', 'mode') as file:
# 文件操作


# 讀取文件內容
with open('C:/path/to/your/example.txt', 'r') as file:
    content = file.read()
    print(content)


# 寫入文件
with open('output.txt', 'w') as file:
    file.write('Hello, world!')
# 追加內容
with open('output.txt', 'a') as file:
    file.write('\nThis is a additional line.')

# 逐行讀取文件
with open('example.txt', 'r') as file:
    for line in file:
        print(line.strip()) # 移除每行末尾的換行符號

# 讀取 TXT 檔案
with open ('file.txt', 'r', encoding='utf-8') as file:
    content = file.read()
    print(content)

# 寫入 TXT 檔案
with open ('output.txt', 'w', encoding='utf-8') as file:
    file.write("這裡是寫入的內容")

